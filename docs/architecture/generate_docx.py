#!/usr/bin/env python3
"""Canvas Artifact Persistence ADR — McKinsey-Quality DOCX Generator"""

import re
import os
from pathlib import Path
from urllib.parse import unquote

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Emu, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.section import WD_SECTION_START
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from PIL import Image as PILImage

# ── Paths ───────────────────────────────────────────────────────────────────
BASE    = Path("/Users/suyash/Desktop/projects/project-context/docs/architecture")
MD_SRC  = BASE / "2026-08-26-canvas-artifact-persistence-adr.md"
OUT     = BASE / "2026-08-26-canvas-artifact-persistence-adr.docx"
IMG_DIR = BASE

DOC_TITLE = "Canvas Artifact Persistence"
HEADER_RIGHT_TEXT = "project-context — Architecture Decision Record"

# ── Colours (hex strings without #) ────────────────────────────────────────
PRIMARY    = "1A4731"
SECONDARY  = "2D6A4F"
ACCENT     = "40916C"
GOLD       = "C9A84C"
LIGHT_BG   = "F0F7F4"
CODE_BG    = "F5F5F5"
C_TEXT     = "1C1C1C"
SUBTLE     = "6B7280"
WHITE      = "FFFFFF"
CALLOUT_BG = "FEFBF0"
DARK_CODE  = "2D2D2D"

def make_rgb(hex6: str) -> RGBColor:
    r, g, b = int(hex6[0:2],16), int(hex6[2:4],16), int(hex6[4:6],16)
    return RGBColor(r, g, b)

# ── XML Helpers ─────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex6: str):
    """Set background fill colour of a table cell."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for existing in tcPr.findall(qn('w:shd')):
        tcPr.remove(existing)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex6.upper())
    tcPr.append(shd)


def set_cell_margins(cell, top_cm=0.1, left_cm=0.2, bottom_cm=0.1, right_cm=0.2):
    """Set internal padding for a table cell."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for existing in tcPr.findall(qn('w:tcMar')):
        tcPr.remove(existing)
    tcMar = OxmlElement('w:tcMar')
    for side, val_cm in [('top', top_cm), ('left', left_cm),
                          ('bottom', bottom_cm), ('right', right_cm)]:
        m = OxmlElement(f'w:{side}')
        twips = int(val_cm / 2.54 * 1440)
        m.set(qn('w:w'),    str(twips))
        m.set(qn('w:type'), 'dxa')
        tcMar.append(m)
    tcPr.append(tcMar)


def remove_table_borders(table):
    """Remove all visible borders from a table."""
    tbl  = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    for existing in tblPr.findall(qn('w:tblBorders')):
        tblPr.remove(existing)
    tblBorders = OxmlElement('w:tblBorders')
    for side in ['top','left','bottom','right','insideH','insideV']:
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'),   'none')
        el.set(qn('w:sz'),    '0')
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), 'auto')
        tblBorders.append(el)
    tblPr.append(tblBorders)


def set_table_borders(table, hex6: str, size_half_pt: int = 4):
    """Apply uniform border to all sides of a table."""
    tbl  = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    for existing in tblPr.findall(qn('w:tblBorders')):
        tblPr.remove(existing)
    tblBorders = OxmlElement('w:tblBorders')
    for side in ['top','left','bottom','right','insideH','insideV']:
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'),   'single')
        el.set(qn('w:sz'),    str(size_half_pt))
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), hex6.upper())
        tblBorders.append(el)
    tblPr.append(tblBorders)


def shade_paragraph(para, hex6: str):
    """Apply background shading to a paragraph."""
    pPr = para._p.get_or_add_pPr()
    for existing in pPr.findall(qn('w:shd')):
        pPr.remove(existing)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex6.upper())
    pPr.append(shd)


def add_run_fmt(run, bold=False, italic=False,
                color_hex=None, size_pt=None, font='Calibri'):
    """Apply formatting to a run."""
    run.bold   = bold
    run.italic = italic
    if font:
        run.font.name = font
    if size_pt:
        run.font.size = Pt(size_pt)
    if color_hex:
        run.font.color.rgb = make_rgb(color_hex)


def add_page_break(doc):
    """Add a page break paragraph."""
    from docx.enum.text import WD_BREAK
    p   = doc.add_paragraph()
    run = p.add_run()
    run.add_break(WD_BREAK.PAGE)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(0)


def set_row_height(row, cm: float):
    """Set exact height for a table row."""
    tr   = row._tr
    trPr = tr.get_or_add_trPr()
    for existing in trPr.findall(qn('w:trHeight')):
        trPr.remove(existing)
    trH = OxmlElement('w:trHeight')
    trH.set(qn('w:val'),   str(int(cm / 2.54 * 1440)))
    trH.set(qn('w:hRule'), 'exact')
    trPr.append(trH)


def set_table_width(table, cm: float):
    """Set table to a fixed width."""
    tbl  = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    for existing in tblPr.findall(qn('w:tblW')):
        tblPr.remove(existing)
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:w'),    str(int(cm / 2.54 * 1440)))
    tblW.set(qn('w:type'), 'dxa')
    tblPr.append(tblW)


def _apply_no_page_break_to_table(table):
    """Repeat header row on overflow pages; let rows flow naturally (no cantSplit)."""
    for i, row in enumerate(table.rows):
        tr   = row._tr
        trPr = tr.get_or_add_trPr()
        for tag in ['w:cantSplit', 'w:tblHeader']:
            for el in trPr.findall(qn(tag)):
                trPr.remove(el)
        if i == 0:
            hdr_el = OxmlElement('w:tblHeader')
            hdr_el.set(qn('w:val'), '1')
            trPr.append(hdr_el)


# ── Document Setup ───────────────────────────────────────────────────────────

def apply_document_defaults(doc):
    """Set A4 page size, margins, default Calibri font, line spacing."""
    section = doc.sections[0]
    section.page_width    = Cm(21.0)
    section.page_height   = Cm(29.7)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.0)
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.header_distance = Cm(1.2)
    section.footer_distance = Cm(1.2)
    section.different_first_page_header_footer = True
    doc.styles['Normal'].font.name = 'Calibri'
    doc.styles['Normal'].font.size = Pt(10.5)
    doc.styles['Normal'].font.color.rgb = make_rgb(C_TEXT)
    nf = doc.styles['Normal'].paragraph_format
    nf.space_after       = Pt(6)
    nf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    nf.line_spacing      = 1.15


def _add_field(run, field_code: str):
    """Insert a Word field (PAGE, NUMPAGES, etc.) into a run."""
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    run._r.append(fldChar1)
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = field_code
    run._r.append(instrText)
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar2)


def setup_header_footer(section, section_name: str = ''):
    """Configure running header and footer for a section."""
    # ── Header ──────────────────────────────────────────────────────────────
    hdr = section.header
    hdr.is_linked_to_previous = False
    for p in hdr.paragraphs:
        p._p.getparent().remove(p._p)

    hdr_para = hdr.add_paragraph()
    hdr_para.paragraph_format.space_before = Pt(0)
    hdr_para.paragraph_format.space_after  = Pt(4)

    r_left = hdr_para.add_run(section_name)
    add_run_fmt(r_left, color_hex=SECONDARY, size_pt=8.5)

    tab = OxmlElement('w:tab')
    r_left._r.append(tab)

    r_right = hdr_para.add_run(HEADER_RIGHT_TEXT)
    add_run_fmt(r_right, color_hex=SUBTLE, size_pt=8.5)
    hdr_para.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Green bottom border on header
    pPr = hdr_para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'),   'single')
    bot.set(qn('w:sz'),    '4')
    bot.set(qn('w:space'), '1')
    bot.set(qn('w:color'), ACCENT.upper())
    pBdr.append(bot)
    pPr.append(pBdr)

    # Tab stop right-aligned at 16.5cm for header
    pPr_h = hdr_para._p.get_or_add_pPr()
    tabs_h = OxmlElement('w:tabs')
    tab_h = OxmlElement('w:tab')
    tab_h.set(qn('w:val'), 'right')
    tab_h.set(qn('w:pos'), str(int(16.5 / 2.54 * 1440)))
    tabs_h.append(tab_h)
    pPr_h.append(tabs_h)

    # ── Footer ──────────────────────────────────────────────────────────────
    ftr = section.footer
    ftr.is_linked_to_previous = False
    for p in ftr.paragraphs:
        p._p.getparent().remove(p._p)

    ftr_para = ftr.add_paragraph()
    ftr_para.paragraph_format.space_before = Pt(4)
    ftr_para.paragraph_format.space_after  = Pt(0)

    r_conf = ftr_para.add_run('Confidential')
    add_run_fmt(r_conf, color_hex=SUBTLE, size_pt=8.5)

    tab2 = OxmlElement('w:tab')
    r_conf._r.append(tab2)
    r_pg = ftr_para.add_run()
    add_run_fmt(r_pg, bold=True, color_hex=C_TEXT, size_pt=8.5)
    _add_field(r_pg, ' PAGE ')

    tab3 = OxmlElement('w:tab')
    r_pg._r.append(tab3)
    r_copy = ftr_para.add_run('project-context — Internal')
    add_run_fmt(r_copy, color_hex=SUBTLE, size_pt=8.5)
    ftr_para.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Tab stops: centre at 8.5cm, right at 16.5cm
    pPr2 = ftr_para._p.get_or_add_pPr()
    tabs = OxmlElement('w:tabs')
    for pos_cm, align in [(8.5, 'center'), (16.5, 'right')]:
        tab_el = OxmlElement('w:tab')
        tab_el.set(qn('w:val'), align)
        tab_el.set(qn('w:pos'), str(int(pos_cm / 2.54 * 1440)))
        tabs.append(tab_el)
    pPr2.append(tabs)

# ── Cover Page ───────────────────────────────────────────────────────────────

def _add_para_to_cell(cell, text, align=WD_ALIGN_PARAGRAPH.CENTER,
                       bold=False, italic=False, color_hex=WHITE,
                       size_pt=12, font='Calibri', space_before_pt=0,
                       space_after_pt=8):
    """Add a styled paragraph to a table cell."""
    para = cell.add_paragraph()
    para.alignment = align
    para.paragraph_format.space_before = Pt(space_before_pt)
    para.paragraph_format.space_after  = Pt(space_after_pt)
    run = para.add_run(text)
    add_run_fmt(run, bold=bold, italic=italic,
                color_hex=color_hex, size_pt=size_pt, font=font)
    return para


def _add_gold_rule_to_cell(cell):
    """Add a centred gold horizontal rule inside a cell via paragraph border."""
    para = cell.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_before = Pt(4)
    para.paragraph_format.space_after  = Pt(4)
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'),   'single')
    bottom.set(qn('w:sz'),    '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), GOLD.upper())
    pBdr.append(bottom)
    pPr.append(pBdr)
    return para


def render_cover(doc):
    """Render a full-page green cover page using a 1x1 table."""
    CONTENT_H_CM = 25.7  # A4 29.7cm - 2cm top - 2cm bottom

    table = doc.add_table(rows=1, cols=1)
    remove_table_borders(table)
    set_table_width(table, 16.5)
    set_row_height(table.rows[0], CONTENT_H_CM)
    cell = table.cell(0, 0)
    set_cell_bg(cell, PRIMARY)
    set_cell_margins(cell, top_cm=2.5, left_cm=2.0, bottom_cm=2.0, right_cm=2.0)

    # Remove the default empty paragraph word adds to the cell
    for p in cell.paragraphs:
        p._p.getparent().remove(p._p)

    # Top spacer
    sp = cell.add_paragraph()
    sp.paragraph_format.space_before = Pt(0)
    sp.paragraph_format.space_after  = Pt(60)

    # Organisation label
    _add_para_to_cell(cell, 'PROJECT-CONTEXT',
                      bold=False, color_hex=GOLD, size_pt=13,
                      space_before_pt=0, space_after_pt=6)

    _add_para_to_cell(cell, 'Agent Platform Engineering',
                      italic=True, color_hex=WHITE, size_pt=10,
                      space_before_pt=0, space_after_pt=40)

    # Main title
    _add_para_to_cell(cell, 'Canvas Artifact Persistence',
                      bold=True, color_hex=WHITE, size_pt=32,
                      space_before_pt=0, space_after_pt=12)

    # Subtitle
    _add_para_to_cell(cell, 'Architecture Decision Record',
                      bold=False, color_hex=GOLD, size_pt=20,
                      space_before_pt=0, space_after_pt=40)

    # Gold separator
    _add_gold_rule_to_cell(cell)

    # Mid spacer
    sp2 = cell.add_paragraph()
    sp2.paragraph_format.space_before = Pt(0)
    sp2.paragraph_format.space_after  = Pt(50)

    # Metadata block
    for label, value in [
        ('Status',         'Accepted'),
        ('Date',           '2026-08-26'),
        ('Prepared by',    'Claude Code'),
        ('Document type',  'Architecture Decision Record'),
        ('Classification', 'Internal'),
    ]:
        meta_para = cell.add_paragraph()
        meta_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        meta_para.paragraph_format.space_before = Pt(0)
        meta_para.paragraph_format.space_after  = Pt(4)
        r_label = meta_para.add_run(f'{label}:  ')
        add_run_fmt(r_label, bold=True, color_hex=GOLD, size_pt=9.5)
        r_value = meta_para.add_run(value)
        add_run_fmt(r_value, color_hex=WHITE, size_pt=9.5)


# ── Section Dividers ─────────────────────────────────────────────────────────

SECTION_SUMMARIES = {
    1: "Canvas output was ephemeral — never stored, never durable, and unverifiable against the agent's own claims.",
    2: "Persist canvas output through the existing upload pipeline, carried on the message's attachments array.",
    3: "Why the artifactRef enum and a new table were both considered and rejected.",
    4: "Four concrete changes: the upload call, a collision-proof key, the tool-loop merge, and the persistence signature.",
    5: "What was actually tested on the write side, the read side, and confirmed by inspection on the frontend.",
    6: "attachments becomes the general carrier for tool-produced files; artifactRef stays reserved for its three entity types.",
}


def _configure_section_margins(sec):
    """Apply standard A4 margins to a Word section."""
    sec.page_width    = Cm(21.0)
    sec.page_height   = Cm(29.7)
    sec.left_margin   = Cm(2.5)
    sec.right_margin  = Cm(2.0)
    sec.top_margin    = Cm(2.0)
    sec.bottom_margin = Cm(2.0)
    sec.header_distance = Cm(1.2)
    sec.footer_distance = Cm(1.2)


def render_section_divider(doc, number: int, title: str):
    """Render a full-page green section divider page (no header/footer)."""
    title_clean = re.sub(r'^SECTION\s+\d+[\s:—–-]*', '', title, flags=re.IGNORECASE).strip()

    # ── Word section for the divider page: suppress header/footer on first page ──
    div_sec = doc.add_section(WD_SECTION_START.NEW_PAGE)
    _configure_section_margins(div_sec)
    div_sec.different_first_page_header_footer = True
    # Clear first-page header and footer so nothing shows on the divider page
    for hf in (div_sec.first_page_header, div_sec.first_page_footer):
        hf.is_linked_to_previous = False
        for p in hf.paragraphs:
            for r in p.runs:
                r.text = ''
    # Regular (non-first-page) header/footer for this section (safety fallback)
    setup_header_footer(div_sec, title_clean)

    CONTENT_H_CM = 25.7
    table = doc.add_table(rows=1, cols=1)
    remove_table_borders(table)
    set_table_width(table, 16.5)
    set_row_height(table.rows[0], CONTENT_H_CM)
    cell = table.cell(0, 0)
    set_cell_bg(cell, PRIMARY)
    set_cell_margins(cell, top_cm=3.0, left_cm=2.5, bottom_cm=2.0, right_cm=2.0)

    for p in cell.paragraphs:
        p._p.getparent().remove(p._p)

    # Large gold section number
    num_para = cell.add_paragraph()
    num_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    num_para.paragraph_format.space_before = Pt(0)
    num_para.paragraph_format.space_after  = Pt(0)
    r_num = num_para.add_run(str(number).zfill(2))
    add_run_fmt(r_num, bold=True, color_hex=GOLD, size_pt=96, font='Calibri')

    # Section title
    title_para = cell.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title_para.paragraph_format.space_before = Pt(0)
    title_para.paragraph_format.space_after  = Pt(12)
    r_title = title_para.add_run(title_clean.upper())
    add_run_fmt(r_title, bold=True, color_hex=WHITE, size_pt=22, font='Calibri')

    # Gold rule
    _add_gold_rule_to_cell(cell)

    # Summary sentence
    summary = SECTION_SUMMARIES.get(number, '')
    if summary:
        sum_para = cell.add_paragraph()
        sum_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        sum_para.paragraph_format.space_before = Pt(10)
        sum_para.paragraph_format.space_after  = Pt(0)
        r_sum = sum_para.add_run(summary)
        add_run_fmt(r_sum, italic=True, color_hex=WHITE, size_pt=11)

    # ── Word section for content pages: restore normal header/footer ──
    content_sec = doc.add_section(WD_SECTION_START.NEW_PAGE)
    _configure_section_margins(content_sec)
    content_sec.different_first_page_header_footer = False
    setup_header_footer(content_sec, title_clean)

# ── Callout Box & Executive Summary ─────────────────────────────────────────

def render_callout_box(doc, headline: str, body_text: str):
    """Render a cream call-out box with gold left border and bold headline."""
    table = doc.add_table(rows=1, cols=1)
    remove_table_borders(table)
    set_table_width(table, 16.0)
    cell = table.cell(0, 0)
    set_cell_bg(cell, CALLOUT_BG)
    set_cell_margins(cell, top_cm=0.2, left_cm=0.4, bottom_cm=0.2, right_cm=0.3)

    # Gold left border only
    tbl   = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    # Remove any existing tblBorders first
    for existing in tblPr.findall(qn('w:tblBorders')):
        tblPr.remove(existing)
    tblBorders = OxmlElement('w:tblBorders')
    left_el = OxmlElement('w:left')
    left_el.set(qn('w:val'),   'single')
    left_el.set(qn('w:sz'),    '24')
    left_el.set(qn('w:space'), '0')
    left_el.set(qn('w:color'), GOLD.upper())
    tblBorders.append(left_el)
    tblPr.append(tblBorders)

    for p in cell.paragraphs:
        p._p.getparent().remove(p._p)

    head_para = cell.add_paragraph()
    head_para.paragraph_format.space_before = Pt(0)
    head_para.paragraph_format.space_after  = Pt(3)
    r_h = head_para.add_run(headline)
    add_run_fmt(r_h, bold=True, color_hex=PRIMARY, size_pt=10.5)

    body_para = cell.add_paragraph()
    body_para.paragraph_format.space_before = Pt(0)
    body_para.paragraph_format.space_after  = Pt(0)
    r_b = body_para.add_run(body_text)
    add_run_fmt(r_b, color_hex=C_TEXT, size_pt=10.0)

    sp = doc.add_paragraph()
    sp.paragraph_format.space_before = Pt(0)
    sp.paragraph_format.space_after  = Pt(8)


def render_exec_summary(doc):
    """Render synthesised 1.5-page Executive Summary."""
    # Heading
    h = doc.add_paragraph()
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    h.paragraph_format.space_before = Pt(0)
    h.paragraph_format.space_after  = Pt(14)
    r_h = h.add_run('Executive Summary')
    add_run_fmt(r_h, bold=True, color_hex=PRIMARY, size_pt=20)

    # Intro paragraph
    intro = (
        "This Technical Proposal presents a production-ready, open-source, cloud-agnostic Unified Service "
        "Delivery Platform for Tea Board India — consolidating thirteen functional modules onto a single, "
        "secure, scalable microservices architecture. The platform is architected to last the full duration "
        "of the contract and beyond: every technology decision is justified against a specific SOW "
        "requirement, every security control is independently verifiable, and every SLA commitment is "
        "backed by a testable architectural guarantee."
    )
    p = doc.add_paragraph(intro)
    p.paragraph_format.space_after = Pt(14)
    for run in p.runs:
        add_run_fmt(run, color_hex=C_TEXT, size_pt=10.5)

    # Five gold call-out highlight boxes
    highlights = [
        (
            "Zero-Trust Security — Seven Independent Layers",
            "Every request is authenticated, authorised, and encrypted at seven independent layers: "
            "Cloudflare DDoS, ModSecurity WAF, Kong API Gateway, Istio mTLS service mesh, Kubernetes "
            "Network Policies, OWASP ASVS L2 application controls, and AES-256-GCM data encryption "
            "with HashiCorp Vault key management. A breach of any single layer does not compromise "
            "the platform."
        ),
        (
            "India Stack Native — Aadhaar, PFMS, DigiLocker, Bhashini, UMANG",
            "The platform integrates natively with UIDAI Aadhaar (AUA e-KYC + Data Vault), PFMS for "
            "DBT disbursement, DigiLocker for certificate delivery, eMudhra eSign, Bhashini for "
            "AI-powered multilingual support in 7 languages, and is pre-designed for UMANG federation "
            "with zero application changes required."
        ),
        (
            "99.9% SLA Backed by Architecture, Not Promises",
            "Multi-AZ Kubernetes deployment, Patroni PostgreSQL HA (< 30-second automatic failover), "
            "Kafka replication factor 3, ArgoCD GitOps for safe zero-downtime deployments, active-passive "
            "DR at CtrlS Mumbai (RTO < 4 hours, RPO < 1 hour), and quarterly DR drills with measured, "
            "documented results submitted to Tea Board India."
        ),
        (
            "13 Modules, One Unified Platform — No Data Silos",
            "All thirteen SOW modules — Website, DBT, eTrading, Green Leaf, Licensing, Grievance, "
            "Tea Mark, MIS/Analytics, Baseline Survey, Inventory, Notification, Returns, and COO — "
            "run on a single unified platform sharing one identity layer (Keycloak), one audit trail "
            "(Kafka + OpenSearch), one notification engine, and one MIS dashboard. Departmental data "
            "isolation is enforced at PostgreSQL Row-Level Security, not at the UI layer."
        ),
        (
            "100% Open-Source — Zero Vendor Lock-in, Zero Licensing Cost",
            "Every component — Spring Boot, Next.js, Flutter, Camunda 8, Kafka, PostgreSQL, "
            "Keycloak, Kong, Istio, ArgoCD, Prometheus, Loki, Jaeger, Wazuh — is open-source with "
            "no proprietary licence dependency. The platform can be migrated between any "
            "MeitY-empanelled cloud provider with only Terraform variable changes."
        ),
    ]
    for headline, body in highlights:
        render_callout_box(doc, headline, body)

    # Spacer
    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # "Platform at a Glance" heading
    tbl_h = doc.add_paragraph()
    tbl_h.paragraph_format.space_before = Pt(10)
    tbl_h.paragraph_format.space_after  = Pt(6)
    r_th = tbl_h.add_run('Platform at a Glance')
    add_run_fmt(r_th, bold=True, color_hex=SECONDARY, size_pt=12)

    rows_data = [
        ('Architecture',   'Three-tier Zero-Trust SOA',                'Defence-in-depth, SOW §8.1.1 compliant'),
        ('Workflow Engine', 'Camunda 8 (Zeebe)',                        '35 BPMN processes as contractual artefacts, SOW §8.1.3'),
        ('Backend',        'Spring Boot 3.2 / Java 17 LTS',            'Official GoI SDK support for UIDAI, PFMS, DGFT'),
        ('Mobile',         'Flutter 3.x (Android + iOS)',               'Offline-first, 3.5 lakh STGs, OWASP MASVS L2'),
        ('Security',       'OWASP ASVS L2 + UIDAI Aadhaar Vault',      'CERT-In VAPT, ISO 27001, DPDPA 2023'),
        ('Infrastructure', 'Kubernetes 1.29 on CtrlS (MeitY)',          'Cloud-agnostic Terraform, IPv6 dual-stack'),
        ('DR',             'CtrlS Hyderabad → Mumbai (active-passive)', 'RTO < 4h, RPO < 1h, quarterly drills'),
        ('Observability',  'Prometheus + Grafana + Loki + Jaeger',      'Real-user monitoring, synthetic checks every 5 min'),
        ('Blockchain',     'Hyperledger Fabric 2.5',                    'Green leaf supply chain, SOW Module 4 §(ix)'),
        ('AI/ML',          'FastAPI + scikit-learn + Prophet',          'DBT forecasting, SOW §9.8.28(7)'),
        ('Compliance',     'DPDPA 2023, GIGW 3.0, IndEA',              'Consent management, IPv6, WCAG 2.1 AA'),
    ]

    table = doc.add_table(rows=len(rows_data) + 1, cols=3)
    set_table_borders(table, ACCENT, size_half_pt=4)
    set_table_width(table, 16.5)
    _apply_no_page_break_to_table(table)

    # Header row
    for ci, hdr_text in enumerate(['Platform Area', 'Technology', 'Key Differentiator']):
        cell = table.cell(0, ci)
        set_cell_bg(cell, PRIMARY)
        set_cell_margins(cell, top_cm=0.1, left_cm=0.2, bottom_cm=0.1, right_cm=0.2)
        for p in cell.paragraphs:
            p._p.getparent().remove(p._p)
        hp = cell.add_paragraph(hdr_text)
        hp.paragraph_format.space_before = Pt(0)
        hp.paragraph_format.space_after  = Pt(0)
        for run in hp.runs:
            add_run_fmt(run, bold=True, color_hex=WHITE, size_pt=9.5)

    # Data rows
    for ri, (area, tech, diff) in enumerate(rows_data):
        bg = LIGHT_BG if ri % 2 == 0 else WHITE
        for ci, text in enumerate([area, tech, diff]):
            cell = table.cell(ri + 1, ci)
            set_cell_bg(cell, bg)
            set_cell_margins(cell, top_cm=0.08, left_cm=0.2, bottom_cm=0.08, right_cm=0.2)
            for p in cell.paragraphs:
                p._p.getparent().remove(p._p)
            dp = cell.add_paragraph(text)
            dp.paragraph_format.space_before = Pt(0)
            dp.paragraph_format.space_after  = Pt(0)
            for run in dp.runs:
                add_run_fmt(run, bold=(ci == 0), color_hex=C_TEXT, size_pt=9.0)

    doc.add_paragraph().paragraph_format.space_after = Pt(4)


# ── Markdown Tokenizer ───────────────────────────────────────────────────────

def tokenize(text: str) -> list:
    """Convert markdown text into a list of typed block tokens."""
    tokens = []
    lines  = text.splitlines()
    i      = 0
    n      = len(lines)

    while i < n:
        line = lines[i]

        # Skip blank lines
        if not line.strip():
            i += 1
            continue

        # Fenced code block ```
        if line.strip().startswith('```'):
            code_lines = []
            i += 1
            while i < n and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i])
                i += 1
            i += 1  # consume closing ```
            tokens.append({'type': 'code_block', 'text': '\n'.join(code_lines)})
            continue

        # HR  ---
        if re.match(r'^-{3,}\s*$', line):
            tokens.append({'type': 'hr'})
            i += 1
            continue

        # Heading  # ## ###
        hm = re.match(r'^(#{1,3})\s+(.*)', line)
        if hm:
            tokens.append({'type': 'heading',
                           'level': len(hm.group(1)),
                           'text':  hm.group(2).strip()})
            i += 1
            continue

        # Image  ![alt](path)
        im = re.match(r'^!\[([^\]]*)\]\(([^)]+)\)\s*$', line.strip())
        if im:
            tokens.append({'type': 'image',
                           'alt':  im.group(1),
                           'path': unquote(im.group(2))})
            i += 1
            continue

        # Pipe table  |...|
        if line.strip().startswith('|'):
            tbl_lines = []
            while i < n and lines[i].strip().startswith('|'):
                tbl_lines.append(lines[i])
                i += 1
            def parse_row(raw):
                return [c.strip() for c in raw.strip().strip('|').split('|')]
            if len(tbl_lines) >= 2:
                headers   = parse_row(tbl_lines[0])
                data_rows = [parse_row(r) for r in tbl_lines[2:]
                             if not re.match(r'^\s*\|?\s*[-:]+\s*\|', r)]
                tokens.append({'type': 'table',
                               'headers': headers,
                               'rows':    data_rows})
            continue

        # Bullet list  - item  or    - sub-item
        if re.match(r'^(\s*)-\s+', line):
            items = []
            while i < n and re.match(r'^(\s*)-\s+', lines[i]):
                m_lvl  = re.match(r'^(\s*)-\s+(.*)', lines[i])
                indent = len(m_lvl.group(1))
                level  = 1 if indent >= 2 else 0
                items.append({'text': m_lvl.group(2).strip(), 'level': level})
                i += 1
            tokens.append({'type': 'bullet_list', 'items': items})
            continue

        # Numbered list  1. item
        if re.match(r'^\d+\.\s+', line):
            items = []
            while i < n and re.match(r'^\d+\.\s+', lines[i]):
                m_nl = re.match(r'^\d+\.\s+(.*)', lines[i])
                items.append(m_nl.group(1).strip())
                i += 1
            tokens.append({'type': 'num_list', 'items': items})
            continue

        # Paragraph (catch-all) — collect consecutive non-special lines
        para_lines = []
        while i < n and lines[i].strip() and \
              not lines[i].strip().startswith('#') and \
              not lines[i].strip().startswith('```') and \
              not lines[i].strip().startswith('|') and \
              not re.match(r'^-{3,}\s*$', lines[i]) and \
              not re.match(r'^(\s*)-\s+', lines[i]) and \
              not re.match(r'^\d+\.\s+', lines[i]) and \
              not re.match(r'^!\[', lines[i].strip()):
            para_lines.append(lines[i])
            i += 1
        if para_lines:
            tokens.append({'type': 'paragraph',
                           'text': ' '.join(l.strip() for l in para_lines)})

    return tokens

# ── Block Renderer ───────────────────────────────────────────────────────────

def render_inline(para, text: str, default_size_pt=10.5,
                  default_color=C_TEXT, default_bold=False):
    """
    Parse inline markdown (**bold**, *italic*, `code`) and add runs to para.
    """
    pattern = r'(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)'
    parts   = re.split(pattern, text)

    for part in parts:
        if not part:
            continue
        if part.startswith('**') and part.endswith('**'):
            run = para.add_run(part[2:-2])
            add_run_fmt(run, bold=True, color_hex=default_color,
                        size_pt=default_size_pt)
        elif part.startswith('*') and part.endswith('*') and len(part) > 2:
            run = para.add_run(part[1:-1])
            add_run_fmt(run, italic=True, color_hex=default_color,
                        size_pt=default_size_pt)
        elif part.startswith('`') and part.endswith('`') and len(part) > 2:
            run = para.add_run(part[1:-1])
            run.font.name  = 'Courier New'
            run.font.bold  = True
            run.font.size  = Pt(default_size_pt - 0.5)
            run.font.color.rgb = make_rgb(PRIMARY)
        else:
            run = para.add_run(part)
            add_run_fmt(run, bold=default_bold, color_hex=default_color,
                        size_pt=default_size_pt)


def render_heading_token(doc, token: dict):
    """Render heading tokens level 1, 2, 3."""
    level = token['level']
    text  = token['text']
    para  = doc.add_paragraph()
    para.paragraph_format.space_after = Pt(6)

    if level == 1:
        para.paragraph_format.space_before = Pt(20)
        run = para.add_run(text)
        add_run_fmt(run, bold=True, color_hex=PRIMARY, size_pt=16)
        pPr = para._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bot = OxmlElement('w:bottom')
        bot.set(qn('w:val'),   'single')
        bot.set(qn('w:sz'),    '6')
        bot.set(qn('w:space'), '1')
        bot.set(qn('w:color'), GOLD.upper())
        pBdr.append(bot)
        pPr.append(pBdr)
    elif level == 2:
        para.paragraph_format.space_before = Pt(14)
        run = para.add_run(text)
        add_run_fmt(run, bold=True, color_hex=SECONDARY, size_pt=13)
    else:
        para.paragraph_format.space_before = Pt(10)
        run = para.add_run(text)
        add_run_fmt(run, bold=True, color_hex=ACCENT, size_pt=11)


def render_hr_token(doc):
    """Suppress --- rules — just add vertical whitespace."""
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(4)
    para.paragraph_format.space_after  = Pt(4)


def render_paragraph_token(doc, token: dict):
    """Render a paragraph token with inline markdown."""
    text = token['text'].strip()
    if not text:
        return
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after  = Pt(6)
    render_inline(para, text)


def render_table_token(doc, token: dict):
    """Render a markdown pipe table with green header and alternating rows."""
    headers   = token['headers']
    data_rows = token['rows']
    if not headers:
        return

    n_cols = len(headers)
    n_rows = len(data_rows)
    table  = doc.add_table(rows=n_rows + 1, cols=n_cols)
    set_table_borders(table, ACCENT, size_half_pt=4)
    set_table_width(table, 16.5)
    _apply_no_page_break_to_table(table)

    # Header row
    for ci, hdr_text in enumerate(headers):
        cell = table.cell(0, ci)
        set_cell_bg(cell, PRIMARY)
        set_cell_margins(cell, top_cm=0.1, left_cm=0.2, bottom_cm=0.1, right_cm=0.2)
        for p in cell.paragraphs:
            p._p.getparent().remove(p._p)
        hp = cell.add_paragraph()
        hp.paragraph_format.space_before = Pt(0)
        hp.paragraph_format.space_after  = Pt(0)
        render_inline(hp, hdr_text, default_size_pt=9.5,
                      default_color=WHITE, default_bold=True)

    # Data rows
    for ri, row_data in enumerate(data_rows):
        bg = LIGHT_BG if ri % 2 == 0 else WHITE
        for ci in range(n_cols):
            cell = table.cell(ri + 1, ci)
            set_cell_bg(cell, bg)
            set_cell_margins(cell, top_cm=0.08, left_cm=0.2,
                             bottom_cm=0.08, right_cm=0.2)
            for p in cell.paragraphs:
                p._p.getparent().remove(p._p)
            dp = cell.add_paragraph()
            dp.paragraph_format.space_before = Pt(0)
            dp.paragraph_format.space_after  = Pt(0)
            cell_text = row_data[ci] if ci < len(row_data) else ''
            render_inline(dp, cell_text, default_size_pt=9.0,
                          default_color=C_TEXT)

    doc.add_paragraph().paragraph_format.space_after = Pt(6)


def render_code_block_token(doc, token: dict):
    """Render structured technical content as clean body-text box with green left bar."""
    table = doc.add_table(rows=1, cols=1)
    remove_table_borders(table)
    set_table_width(table, 16.0)
    cell = table.cell(0, 0)
    set_cell_bg(cell, LIGHT_BG)
    set_cell_margins(cell, top_cm=0.25, left_cm=0.5, bottom_cm=0.25, right_cm=0.4)

    # Thin ACCENT green left border — clean, professional
    tbl   = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    for existing in tblPr.findall(qn('w:tblBorders')):
        tblPr.remove(existing)
    tblBorders = OxmlElement('w:tblBorders')
    left_el = OxmlElement('w:left')
    left_el.set(qn('w:val'),   'single')
    left_el.set(qn('w:sz'),    '12')
    left_el.set(qn('w:space'), '0')
    left_el.set(qn('w:color'), ACCENT.upper())
    tblBorders.append(left_el)
    tblPr.append(tblBorders)

    for p in cell.paragraphs:
        p._p.getparent().remove(p._p)

    for line in token['text'].split('\n'):
        display = line if line.strip() else ''
        # Detect indentation level for visual hierarchy
        stripped  = line.lstrip()
        n_spaces  = len(line) - len(stripped)
        indent_cm = min(n_spaces * 0.18, 2.5)

        cp = cell.add_paragraph(display if display else '')
        cp.paragraph_format.space_before    = Pt(0)
        cp.paragraph_format.space_after     = Pt(1)
        cp.paragraph_format.left_indent     = Cm(indent_cm)
        for run in cp.runs:
            run.font.name      = 'Calibri'
            run.font.size      = Pt(9.5)
            run.font.color.rgb = make_rgb(C_TEXT)
            # Bold the "Stage N:" / "Step N:" / top-level labels
            first_word = stripped.split()[0] if stripped.split() else ''
            if (n_spaces == 0 and stripped and
                    not stripped.startswith(('→', '←', '├', '└', '|', '-'))):
                run.font.bold = True

    doc.add_paragraph().paragraph_format.space_after = Pt(6)


def render_image_token(doc, token: dict):
    """Render an image centred with gold rule above and italic caption below."""
    alt      = token['alt']
    rel_path = token['path']
    img_path = BASE / rel_path

    if not img_path.exists():
        render_callout_box(doc,
            f'[DIAGRAM — {img_path.name}]',
            f'Image file not found: {img_path}\nExport the diagram and re-run.')
        return

    try:
        with PILImage.open(img_path) as pil_img:
            w_px, h_px = pil_img.size
            dpi_info   = pil_img.info.get('dpi', (96, 96))
            dpi        = dpi_info[0] if dpi_info and dpi_info[0] else 96
        w_cm       = w_px / dpi * 2.54
        display_w  = Cm(min(w_cm, 14.0))
    except Exception:
        display_w = Cm(14.0)

    # Gold rule above
    rule_para = doc.add_paragraph()
    rule_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rule_para.paragraph_format.space_before = Pt(4)
    rule_para.paragraph_format.space_after  = Pt(2)
    pPr = rule_para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    top_el = OxmlElement('w:top')
    top_el.set(qn('w:val'),   'single')
    top_el.set(qn('w:sz'),    '4')
    top_el.set(qn('w:space'), '1')
    top_el.set(qn('w:color'), GOLD.upper())
    pBdr.append(top_el)
    pPr.append(pBdr)

    # Image paragraph
    img_para = doc.add_paragraph()
    img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    img_para.paragraph_format.space_before = Pt(0)
    img_para.paragraph_format.space_after  = Pt(6)
    run = img_para.add_run()
    run.add_picture(str(img_path), width=display_w)

    # Caption
    cap_para = doc.add_paragraph()
    cap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap_para.paragraph_format.space_before = Pt(0)
    cap_para.paragraph_format.space_after  = Pt(6)
    r_cap = cap_para.add_run(f'Figure: {alt}')
    add_run_fmt(r_cap, italic=True, color_hex=SUBTLE, size_pt=9.0)


def render_bullet_list_token(doc, token: dict):
    """Render bullet list: ▪ first-level, – second-level."""
    for item in token['items']:
        para = doc.add_paragraph()
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after  = Pt(3)
        if item['level'] == 0:
            para.paragraph_format.left_indent       = Cm(0.5)
            para.paragraph_format.first_line_indent = Cm(-0.35)
            bullet_run = para.add_run('▪  ')
            add_run_fmt(bullet_run, bold=True, color_hex=ACCENT, size_pt=10.5)
            render_inline(para, item['text'])
        else:
            para.paragraph_format.left_indent       = Cm(1.2)
            para.paragraph_format.first_line_indent = Cm(-0.35)
            bullet_run = para.add_run('–  ')
            add_run_fmt(bullet_run, color_hex=SUBTLE, size_pt=10.0)
            render_inline(para, item['text'], default_size_pt=10.0,
                          default_color=C_TEXT)


def render_num_list_token(doc, token: dict):
    """Render a numbered list."""
    for idx, item_text in enumerate(token['items'], 1):
        para = doc.add_paragraph()
        para.paragraph_format.space_before      = Pt(0)
        para.paragraph_format.space_after       = Pt(3)
        para.paragraph_format.left_indent       = Cm(0.6)
        para.paragraph_format.first_line_indent = Cm(-0.5)
        num_run = para.add_run(f'{idx}.  ')
        add_run_fmt(num_run, bold=True, color_hex=ACCENT, size_pt=10.5)
        render_inline(para, item_text)


def render_block(doc, token: dict):
    """Dispatch a single token to its renderer."""
    t = token['type']
    if   t == 'heading':     render_heading_token(doc, token)
    elif t == 'table':       render_table_token(doc, token)
    elif t == 'code_block':  render_code_block_token(doc, token)
    elif t == 'image':       render_image_token(doc, token)
    elif t == 'bullet_list': render_bullet_list_token(doc, token)
    elif t == 'num_list':    render_num_list_token(doc, token)
    elif t == 'hr':          render_hr_token(doc)
    elif t == 'paragraph':   render_paragraph_token(doc, token)


def _estimate_page_numbers(tokens: list) -> dict:
    """
    Pre-pass: estimate which page each H1/H2 heading lands on.
    Returns {heading_text: page_number}.
    """
    LINES = 50  # body lines per A4 page at 10.5pt/1.15 spacing

    def add(n):
        nonlocal page, line
        line += n
        while line >= LINES:
            page += 1
            line -= LINES

    page_map = {}
    page = 1   # cover
    line = 0

    # TOC occupies 1 page for a doc this short
    toc_count = sum(1 for t in tokens if t['type'] == 'heading' and t['level'] <= 2)
    add(toc_count * 1.2 + 4)

    for tok in tokens:
        t = tok['type']
        if t == 'heading':
            lvl  = tok['level']
            text = tok['text']
            if lvl == 1:
                # Each H1 section starts on a new page after its divider page
                page += 1        # divider page
                line  = 0
                page += 1        # content page start
                line  = 0
                page_map[text] = page
                add(3)
            elif lvl == 2:
                page_map[text] = page
                add(2.5)
            else:
                add(2)
        elif t == 'paragraph':
            chars = len(tok.get('text', ''))
            add(max(1, chars // 85) + 0.5)
        elif t == 'table':
            add(len(tok.get('rows', [])) * 1.4 + 2)
        elif t == 'code_block':
            add(min(tok['text'].count('\n') + 1, 35) * 0.9 + 1)
        elif t == 'bullet_list':
            add(len(tok.get('items', [])) * 1.1 + 0.5)
        elif t == 'num_list':
            add(len(tok.get('items', [])) * 1.1 + 0.5)
        elif t == 'image':
            add(12)
        elif t == 'hr':
            add(0.5)

    return page_map


def render_toc(doc, tokens: list):
    """Render Table of Contents with estimated page numbers and dot leaders."""
    page_map = _estimate_page_numbers(tokens)

    toc_h = doc.add_paragraph()
    toc_h.paragraph_format.space_before = Pt(0)
    toc_h.paragraph_format.space_after  = Pt(18)
    r_toc = toc_h.add_run('Contents')
    add_run_fmt(r_toc, bold=True, color_hex=PRIMARY, size_pt=20)

    TAB_POS = str(int(15.5 / 2.54 * 1440))  # 15.5 cm right-aligned

    entries = [(t['level'], t['text'])
               for t in tokens
               if t['type'] == 'heading' and t['level'] <= 2]

    for level, text in entries:
        pg = page_map.get(text, '')
        ep = doc.add_paragraph()
        ep.paragraph_format.space_before = Pt(0)
        ep.paragraph_format.space_after  = Pt(2)

        # Tab stop with dot leader at right margin
        pPr  = ep._p.get_or_add_pPr()
        tabs = OxmlElement('w:tabs')
        te   = OxmlElement('w:tab')
        te.set(qn('w:val'),    'right')
        te.set(qn('w:pos'),    TAB_POS)
        te.set(qn('w:leader'), 'dot')
        tabs.append(te)
        pPr.append(tabs)

        if level == 1:
            ep.paragraph_format.left_indent = Cm(0)
            ep.paragraph_format.space_before = Pt(6)
            r = ep.add_run(text)
            add_run_fmt(r, bold=True, color_hex=PRIMARY, size_pt=10.5)
        else:
            ep.paragraph_format.left_indent = Cm(0.6)
            r = ep.add_run(text)
            add_run_fmt(r, color_hex=SECONDARY, size_pt=9.5)

        # Tab → page number
        tab_run = ep.add_run()
        tab_el  = OxmlElement('w:tab')
        tab_run._r.append(tab_el)
        pg_run  = ep.add_run(str(pg) if pg else '')
        if level == 1:
            add_run_fmt(pg_run, bold=True, color_hex=PRIMARY, size_pt=10.5)
        else:
            add_run_fmt(pg_run, color_hex=SECONDARY, size_pt=9.5)


# ── Main Orchestrator ────────────────────────────────────────────────────────

SECTION_HEADERS = {
    1: 'Section 1 — Context',
    2: 'Section 2 — Decision',
    3: 'Section 3 — Alternatives Considered',
    4: 'Section 4 — Implementation',
    5: 'Section 5 — Verification',
    6: 'Section 6 — Consequences',
}


def split_into_sections(tokens: list) -> list:
    """
    Group tokens by top-level section heading.
    Returns list of dicts: {'number': int, 'title': str, 'tokens': list}
    """
    sections = []
    current  = None
    sec_re   = re.compile(r'SECTION\s+(\d+)', re.IGNORECASE)

    for token in tokens:
        if token['type'] == 'heading' and token['level'] == 1:
            m = sec_re.search(token['text'])
            if m:
                if current:
                    sections.append(current)
                current = {
                    'number': int(m.group(1)),
                    'title':  token['text'],
                    'tokens': [token],
                }
                continue
        if current is not None:
            current['tokens'].append(token)

    if current:
        sections.append(current)

    return sections


def main():
    print(f'Reading {MD_SRC} …')
    text   = MD_SRC.read_text(encoding='utf-8')
    tokens = tokenize(text)
    print(f'  Tokenized: {len(tokens)} blocks')

    sections = split_into_sections(tokens)
    print(f'  Found {len(sections)} sections')

    doc = Document()
    apply_document_defaults(doc)
    setup_header_footer(doc.sections[0], DOC_TITLE)

    # Cover
    print('Rendering cover page …')
    render_cover(doc)
    add_page_break(doc)   # cover → TOC

    # Table of Contents
    print('Rendering table of contents …')
    render_toc(doc, tokens)
    add_page_break(doc)   # TOC → exec summary

    # No Executive Summary — this is a short, single-decision ADR; the
    # proposal-style highlight-box summary doesn't fit that shape.
    # No page break here — render_section_divider uses a NEXT_PAGE section break

    # Sections
    for sec in sections:
        num       = sec['number']
        title     = sec['title']
        print(f'  Section {num}: {title[:55]} …')

        render_section_divider(doc, num, title)

        # Skip the H1 heading — it's already on the divider page
        body_tokens = [
            t for t in sec['tokens']
            if not (t['type'] == 'heading' and t['level'] == 1
                    and re.search(r'SECTION\s+\d+', t['text'], re.IGNORECASE))
        ]
        for token in body_tokens:
            render_block(doc, token)

    # Save
    print(f'\nSaving to {OUT} …')
    doc.save(str(OUT))

    print('\n✓ Done.')
    print(f'  Paragraphs : {len(doc.paragraphs)}')
    print(f'  Tables     : {len(doc.tables)}')
    print(f'  Output     : {OUT}')


if __name__ == '__main__':
    main()
